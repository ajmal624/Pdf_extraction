# app.py
import streamlit as st
import pandas as pd
from docx import Document
import io
import re

st.set_page_config(layout="wide")

# ---------- Helpers ----------
def normalize_text(text: str) -> str:
    if text is None:
        return ""
    # collapse whitespace/newlines/tabs, strip quotes
    text = re.sub(r'[\r\n\t]+', ' ', text)
    text = re.sub(r'\s+', ' ', text).strip()
    text = text.strip(' "\'`')  # trim surrounding quotes
    return text

def is_probable_section_title(text: str) -> bool:
    """
    Heuristics to decide whether a short cell/paragraph is a section header/title:
      - contains words like 'information', 'details', 'section', 'form', 'info'
      - OR the row has exactly one non-empty cell (we treat these single-cell rows as section rows)
    This intentionally errs on the side of skipping generic section labels, not real fields like "Date".
    """
    if not text:
        return False
    t = text.lower()
    keywords = ['information', 'info', 'details', 'section', 'form', 'appraisal']
    if any(k in t for k in keywords):
        return True
    # also treat lines that look like big titles: more than 1 word but short and ends with ':'? (we treat those as not section)
    return False

def extract_from_table(table):
    """
    Walk through a table row-by-row and extract (field, value) pairs.
    Handles rows where many cells are present and multiple field/value pairs appear in the same row.
    """
    pairs = []
    for row in table.rows:
        cells = [normalize_text(c.text) for c in row.cells]
        non_empty = [c for c in cells if c]
        if not non_empty:
            continue

        # If the row only contains a single non-empty cell it's likely a section header -> skip
        if len(non_empty) == 1 and is_probable_section_title(non_empty[0]):
            continue

        # Now parse the sequence of non-empty items left-to-right into (field, value) pairs:
        i = 0
        while i < len(non_empty):
            text = non_empty[i]

            # skip probable section titles if they appear in-line
            if is_probable_section_title(text):
                i += 1
                continue

            # take the current cell as field, next cell as value (if exists)
            field = text
            value = non_empty[i + 1] if (i + 1) < len(non_empty) else ""
            # If value itself looks like a section title (rare), skip the pair and continue
            if is_probable_section_title(value) and not value:
                i += 1
                continue

            pairs.append((field, value))
            i += 2  # move to next possible pair
    return pairs

def extract_from_paragraphs(paragraphs):
    """
    Extract field:value from paragraphs using colon split.
    Skip short lines that look like section titles (no colon and short).
    """
    pairs = []
    for p in paragraphs:
        line = normalize_text(p)
        if not line:
            continue
        # If it has a colon, treat as field:value
        if ":" in line:
            field, value = line.split(":", 1)
            field, value = normalize_text(field), normalize_text(value)
            # skip section-looking fields
            if field and not is_probable_section_title(field):
                pairs.append((field, value))
        else:
            # skip short lines likely to be section headers; keep longer lines only if they look like a value for a previous field
            words = line.split()
            if len(words) > 6:
                # treat as possible unlabeled value (rare). We skip to avoid 'Unlabeled' columns.
                # If you'd like these attached to last field, we could implement that behavior.
                pass
    return pairs

# ---------- Main extraction function ----------
def extract_docx_auto(docx_file):
    doc = Document(docx_file)
    # collect pairs
    collected = []

    # 1) Tables (first, because forms often use tables)
    for table in doc.tables:
        collected.extend(extract_from_table(table))

    # 2) Paragraphs
    paragraph_texts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    collected.extend(extract_from_paragraphs(paragraph_texts))

    # Normalize keys and keep last-occurrence for duplicates (typical expectation)
    cleaned = []
    for k, v in collected:
        key = normalize_text(k)
        val = normalize_text(v)
        # ignore obvious section headings accidentally left
        if not key:
            continue
        if is_probable_section_title(key):
            continue
        cleaned.append((key, val))

    # dedupe keeping last value
    ordered = {}
    for k, v in cleaned:
        ordered[k] = v

    # return as list of pairs (for editing) and as single-row dict
    pairs = list(ordered.items())
    row_dict = ordered  # field -> value
    return pairs, row_dict

# ---------- Streamlit UI ----------
st.title("DOCX → Clean CSV (auto-detect section headers)")

uploaded = st.file_uploader("Upload a DOCX file", type=["docx"])
if not uploaded:
    st.info("Upload a .docx file and I'll try to extract fields and values into a clean CSV.")
    st.stop()

pairs, row_dict = extract_docx_auto(uploaded)

st.markdown("### Detected Field → Value pairs (edit or remove rows if needed)")
if pairs:
    pairs_df = pd.DataFrame(pairs, columns=["Field", "Value"])
else:
    pairs_df = pd.DataFrame(columns=["Field", "Value"])

# allow user to edit the detected pairs
edited = st.experimental_data_editor(pairs_df, num_rows="dynamic")

# Button to build final single-row table
if st.button("Build final table and download CSV"):
    # Validate: drop empty fields
    edited = edited.dropna(subset=["Field"]).copy()
    edited["Field"] = edited["Field"].astype(str).map(normalize_text)
    edited["Value"] = edited["Value"].astype(str).map(normalize_text)

    # Build a single-row DataFrame where columns are fields
    final_dict = {}
    for _, r in edited.iterrows():
        k = r["Field"]
        v = r["Value"]
        if k:
            final_dict[k] = v

    if not final_dict:
        st.error("No valid field/value pairs found. Try editing the table to add fields, or upload another file.")
    else:
        final_df = pd.DataFrame([final_dict])
        st.markdown("### Final single-row table (fields are columns)")
        st.dataframe(final_df)

        # prepare CSV
        buffer = io.StringIO()
        final_df.to_csv(buffer, index=False)
        b = buffer.getvalue().encode("utf-8")

        st.download_button(
            "Download CSV",
            data=b,
            file_name="extracted_doc.csv",
            mime="text/csv"
        )

# show a helpful note about heuristics
st.markdown(
    """
    **Notes on heuristics**
    - The app skips probable section headers (e.g. "Client Information", "Property Information") using keyword heuristics rather than a fixed list.
    - Table rows that contain multiple (field, value) pairs are parsed left-to-right.  
    - If the automatic parsing still creates noisy columns, edit the pairs above and then click **Build final table and download CSV**.
    - If you'd like automatic attachment of orphan lines (multi-line values) to the previous field, I can add that behavior — say so and I'll provide an updated version.
    """
)
